import sys
import sqlite3 as sql
import os
import csv
from datetime import date
from sqlite3 import Error
from openpyxl import Workbook, load_workbook
from num2words import num2words

from docx import Document
from docx.shared import Inches

def printExcel():
  try:

    # Connect to database
    conn=sql.connect('database.db')

  # Export data into CSV file
    print("Exporting data into CSV............")
    cursor = conn.cursor()
    cursor.execute("select * from User where monitorando = False")
    with open("checklist.csv", "w") as csv_file:
        csv_writer = csv.writer(csv_file, delimiter="\t")
        csv_writer.writerow([i[0] for i in cursor.description])
        csv_writer.writerows(cursor)

    dirpath = os.getcwd() + "/employee_data.csv"
    print("Data exported Successfully into {}".format(dirpath))

  except Error as e:
    print(e)

  # Close database connection
  finally:
    conn.close()

def impContrato(nome,adress,city,cpf,kWp,placa,inver,preco):
  document = Document()
  document.add_picture('ecollislogo.png', width=Inches(1.25))
  document.add_heading('CONTRATO DE INSTALAÇÃO', level=1)

  p = document.add_paragraph('Na data de ')
  p.add_run(str(date.today()))
  p.add_run(', a ')
  p.add_run(str(nome))
  p.add_run(', pessoa jurídica de direito privado, inscrita sob o CPF/CNPJ ')
  p.add_run(str(cpf))
  p.add_run(', situado no endereço ')
  p.add_run(str(adress))
  p.add_run(' na cidade de ')
  p.add_run(str(city))
  p.add_run('. E a Empresa ECOLLIS ENERGIA SOLAR, estabelecida na Cidade de Paranavaí, Pr, sito à Rua Rio Grande do Norte, Nº1970, inscrita no CNPJ sob nº 29.737.462/0001-78, neste ato representada pelo Sr.(a) CLEBER PEREIRA MARQUES, Brasileiro, Casado, inscrito sob o CPF nº 005.075.819-54, doravante denominada CONTRATADO, resolvem celebrar o presente Contrato de Prestação de Serviços, mediante as seguintes cláusulas e condições. ')
  
  
  document.add_heading('Clausula Primeira',level=2).bold = True
  h = document.add_paragraph('Do Objetivo: \n')
  h.add_run('O presente contrato tem por objeto a prestação de serviços técnicos especializados junto a empresa Contratada.\n\n')
  
  h.add_run('§1º ').bold = True
  h.add_run(' A realização dos serviços ocorrerá conforme necessidade da empresa dentro da técnica de Instalação Especializada.\n\n')
  
  h.add_run('§2º ').bold = True
  h.add_run('Sistema gerador fotovoltaico com potência de ')
  h.add_run(str(kWp))
  h.add_run('kWp, composto por: ')
  h.add_run(str(placa))
  h.add_run('; ')
  h.add_run(str(inver))
  h.add_run('; Estrutura de fixação para telhado; Cabos; String box; Conectores; Projeto/Homologação e Serviço de instalação.')
  
  document.add_heading('Clausula Segunda',level=2).bold = True
  h = document.add_paragraph('Das obrigações das partes:\n')
  h.add_run('As partes se obrigam a cumprir fielmente o presente contrato nos termos a seguir: \n')
  
  h.add_run('§1º ').bold = True
  h.add_run('Constituem obrigações do')
  h.add_run(' CONTRATADO: ').bold = True
  h.add_run('Cumprir o presente contrato prestando os serviços de INSTALAÇÃO DE SISTEMA GERADOR FOTOVOLTAICO dentro da necessidade do ')
  h.add_run('CONTRATANTE')
  h.add_run(' para geração de energia elétrica como micro gerador com configuração de sistema ONGRID (conectado à rede elétrica da concessionária).\n')
  h.add_run('O')
  h.add_run(' CONTRATADO ').bold = True
  h.add_run('se compromete em realizar inspeção e manutenção de seis em seis meses, no período de 1 ano, para o perfeito funcionamento do sistema comprometendo-se a fornecer garantia de instalação do sistema de 1 ano.\n\n')
  
  h.add_run('§2º ').bold = True
  h.add_run('Constituem obrigações da ')
  h.add_run('CONTRATANTE: ').bold = True
  h.add_run('Colocar à disposição do')
  h.add_run(' CONTRATADO ').bold = True
  h.add_run('todas as informações necessárias para realizar seu trabalho.\n')

  document.add_heading('Clausula Terceira',level=2).bold = True
  h = document.add_paragraph('Preço:\n')
  h.add_run('R$ ')
  dinheiro = '{:,.2f}'.format(int(preco))
  h.add_run(dinheiro)
  h.add_run(' (')
  h.add_run(number_to_long_number(preco))
  h.add_run(')')

  document.add_heading('Clausula Quarta',level=2).bold = True
  h = document.add_paragraph('Condições de Pagamento:\n')
  h.add_run('Através de financiamento\n')
  h.add_run("O prazo de vigência do presente contrato é de 12 (doze) meses, podendo ser prorrogado por igual ou menor prazo, se as partes assim concordarem.\n")
  
  document.add_heading('Clausula Quinta',level=2).bold = True
  h = document.add_paragraph('Este Contrato será rescindido automaticamente ao final da sua vigência, tornando-se vencido e, assim, executável, independente de manifestação das partes se o CONTRATANTE deixar de efetuar o pagamento de acordo com a cláusula terceira.\n\n')
  h.add_run('§1º ').bold = True
  h.add_run('Se o CONTRATANTE rescindir justificadamente por não aprovação de crédito bancário, o presente contrato perderá sua validade não onerando encargos ao CONTRATANTE.\n\n')
  h.add_run('§2º ').bold = True
  h.add_run('Se o CONTRATADO rescindir injustificadamente o presente contrato sem concluir integralmente todas as fases do presente projeto, perderá todos os direitos autorais sobre as fases já concluídas, sub-rogando tais direitos a qualquer outro profissional que vier a ser contratado pelo CONTRATANTE.\n')
  
  document.add_heading('Clausula Sexta',level=2).bold = True
  h=document.add_paragraph('Fica eleito o foro da Comarca de Paranavaí, Pr para dirimir quaisquer dúvidas da aplicação do presente contrato.\n\n')
   
  h.add_run(' Paranavaí, Paraná')
  h.add_run(formated(date.today()))

  h.add_run('\n\n\n\n\n _______________________________                                                       _______________________________')
  h.add_run('                                              CONTRATANTE                                                                                CONTRATADO\n\n\n\n\n\n\n\n')
  h.add_run('\n _______________________________                                                       _______________________________')
  h.add_run('                                              TESTEMUNHA                                                                                TESTEMUNHA')
  

  document.save('contrato.docx')


def number_to_long_number(number_p):
    if number_p.find(',')!=-1:
        number_p = number_p.split(',')
        number_p1 = int(number_p[0].replace('.',''))
        number_p2 = int(number_p[1])
    else:
        number_p1 = int(number_p.replace('.',''))
        number_p2 = 0    
        
    if number_p1 == 1:
        aux1 = ' real'
    else:
        aux1 = ' reais'
        
    if number_p2 == 1:
        aux2 = ' centavo'
    else:
        aux2 = ' centavos'
        
    text1 = ''
    if number_p1 > 0:
        text1 = num2words(number_p1,lang='pt_BR') + str(aux1)
    else:
        text1 = ''
    
    if number_p2 > 0:
        text2 = num2words(number_p2,lang='pt_BR') + str(aux2) 
    else: 
        text2 = ''
    
    if (number_p1 > 0 and number_p2 > 0):
        result = text1 + ' e ' + text2
    else:
        result = text1 + text2

    return result



def formated(self):
    month_name = {
        '1': 'janeiro',
        '2': 'fevereiro',
        '3': 'março',
        '4': 'abril',
        '5': 'maio',
        '6': 'junho',
        '7': 'julho',
        '8': 'agosto',
        '9': 'setembro',
        '10': 'outubro',
        '11': 'novembro',
        '12': 'dezembro'        
        }
    numeric = f'{self.day:02d}/{self.month:02d}/{self.year:04d}'
    written = f'{str(self.day)} de {month_name[str(self.month)]} de {str(self.year)}'
    ret = numeric + ' - ' + written
    print(ret)
    return ret